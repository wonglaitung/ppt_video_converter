#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF 转 Word 工具模块
将 PDF 每页转为图片并插入 Word 文档

功能：
1. 优先使用 PyMuPDF (fitz) 将 PDF 每页转换为图片（纯 Python，无需系统依赖）
2. 备选使用 poppler-utils (pdftoppm) 转换
3. 使用 python-docx 将图片逐页插入 Word 文档
"""

import os
import sys
import argparse
import tempfile
import subprocess
import shutil
from pathlib import Path
from datetime import datetime

# 添加项目根目录到路径
script_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.dirname(script_dir)
if project_root not in sys.path:
    sys.path.insert(0, project_root)


# ============================================================================
# 依赖检查
# ============================================================================

def check_poppler_available():
    """检查 poppler-utils 是否可用"""
    try:
        result = subprocess.run(
            ['pdftoppm', '-v'],
            capture_output=True,
            timeout=5
        )
        return result.returncode == 0
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False


def check_python_docx_available():
    """检查 python-docx 是否可用"""
    try:
        from docx import Document
        return True
    except ImportError:
        return False


def check_pymupdf_available():
    """检查 PyMuPDF 是否可用"""
    try:
        import fitz
        return True
    except ImportError:
        return False


# ============================================================================
# PDF 转图片
# ============================================================================

def convert_pdf_to_images_pymupdf(pdf_path, output_dir, dpi=150, quality=75, use_jpeg=True):
    """使用 PyMuPDF 将 PDF 转换为图片序列

    Args:
        pdf_path (str): PDF 文件路径
        output_dir (str): 输出目录
        dpi (int): DPI 值（默认 150）
        quality (int): JPEG 质量 1-100（默认 75）
        use_jpeg (bool): 是否使用 JPEG 格式（默认 True，压缩率更高）

    Returns:
        list: 图片文件路径列表（按页码排序）
    """
    import fitz

    os.makedirs(output_dir, exist_ok=True)

    # 打开 PDF
    doc = fitz.open(pdf_path)
    image_files = []

    # 计算缩放因子（DPI 72 是 PDF 默认）
    zoom = dpi / 72
    matrix = fitz.Matrix(zoom, zoom)

    img_ext = 'jpg' if use_jpeg else 'png'

    for page_num in range(len(doc)):
        page = doc[page_num]

        # 渲染页面为图片
        pix = page.get_pixmap(matrix=matrix)

        # 保存图片
        img_path = os.path.join(output_dir, f"page_{page_num + 1:04d}.{img_ext}")
        if use_jpeg:
            pix.save(img_path, output="jpeg", jpg_quality=quality)
        else:
            pix.save(img_path)
        image_files.append(img_path)

    doc.close()
    return image_files


def convert_pdf_to_images_poppler(pdf_path, output_dir, dpi=150):
    """使用 poppler-utils (pdftoppm) 将 PDF 转换为图片序列

    Args:
        pdf_path (str): PDF 文件路径
        output_dir (str): 输出目录
        dpi (int): DPI 值（默认 150）

    Returns:
        list: 图片文件路径列表（按页码排序）
    """
    os.makedirs(output_dir, exist_ok=True)

    # 使用 pdftoppm 转换 PDF 为 PNG 图片
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]

    cmd = [
        'pdftoppm',
        '-png',
        '-r', str(dpi),
        pdf_path,
        os.path.join(output_dir, base_name)
    ]

    result = subprocess.run(cmd, capture_output=True, text=True)

    if result.returncode != 0:
        raise RuntimeError(f"PDF 转换失败: {result.stderr}")

    # 收集生成的图片文件（按页码排序）
    image_files = []
    for file in sorted(os.listdir(output_dir)):
        if file.startswith(base_name) and file.endswith('.png'):
            image_files.append(os.path.join(output_dir, file))

    if not image_files:
        raise RuntimeError("未生成任何图片文件")

    return image_files


def convert_pdf_to_images(pdf_path, output_dir, dpi=150, quality=75, use_jpeg=True):
    """将 PDF 转换为图片序列（自动选择可用方案）

    优先级：PyMuPDF > poppler-utils

    Args:
        pdf_path (str): PDF 文件路径
        output_dir (str): 输出目录
        dpi (int): DPI 值（默认 150）
        quality (int): JPEG 质量 1-100（默认 75）
        use_jpeg (bool): 是否使用 JPEG 格式（默认 True）

    Returns:
        list: 图片文件路径列表（按页码排序）
    """
    # 优先使用 PyMuPDF
    if check_pymupdf_available():
        print("   📦 使用 PyMuPDF 转换")
        return convert_pdf_to_images_pymupdf(pdf_path, output_dir, dpi, quality, use_jpeg)

    # 备选使用 poppler-utils
    if check_poppler_available():
        print("   📦 使用 poppler-utils 转换")
        return convert_pdf_to_images_poppler(pdf_path, output_dir, dpi)

    raise RuntimeError(
        "无可用的 PDF 转换工具。\n"
        "请安装其中之一：\n"
        "  pip install PyMuPDF  # 推荐，无需系统依赖\n"
        "  sudo apt-get install poppler-utils  # 系统工具"
    )


# ============================================================================
# Word 文档生成
# ============================================================================

def create_word_with_images(image_files, output_path, title=None):
    """创建包含图片的 Word 文档

    Args:
        image_files (list): 图片文件路径列表
        output_path (str): 输出 Word 文件路径
        title (str): 文档标题（可选）

    Returns:
        str: Word 文件路径
    """
    from docx import Document
    from docx.shared import Inches, Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # 创建文档
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # 添加标题（可选）
    if title:
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.bold = True
        title_run.font.size = Pt(18)
        doc.add_paragraph()  # 空行

    # 遍历图片，每页插入一张
    for i, img_path in enumerate(image_files):
        # 添加段落并插入图片
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()

        # 插入图片，设置宽度为 A4 页面宽度（约 7.5 英寸）
        run.add_picture(img_path, width=Inches(7.5))

        # 如果不是最后一页，添加分页符
        if i < len(image_files) - 1:
            doc.add_page_break()

    # 保存文档
    doc.save(output_path)

    return output_path


# ============================================================================
# 主转换函数
# ============================================================================

def pdf_to_word_with_images(pdf_path, output_path=None, dpi=100, quality=70, keep_temp=False):
    """将 PDF 每页转为图片并插入 Word 文档

    Args:
        pdf_path (str): PDF 文件路径
        output_path (str): 输出 Word 文件路径（可选）
        dpi (int): 图片 DPI（默认 100，压缩后适合阅读）
        quality (int): JPEG 质量 1-100（默认 70）
        keep_temp (bool): 是否保留临时图片文件

    Returns:
        dict: {'success': True, 'word_path': '...', 'page_count': N}
    """
    print("=" * 50)
    print("📄 PDF 转 Word 工具")
    print("=" * 50)
    print(f"📁 输入文件: {pdf_path}")
    print()

    # 检查输入文件
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF 文件不存在: {pdf_path}")

    # 检查依赖
    if not check_pymupdf_available() and not check_poppler_available():
        raise RuntimeError(
            "无可用的 PDF 转换工具。\n"
            "请安装其中之一：\n"
            "  pip install PyMuPDF  # 推荐，无需系统依赖\n"
            "  sudo apt-get install poppler-utils  # 系统工具"
        )

    if not check_python_docx_available():
        raise RuntimeError("python-docx 未安装，请运行: pip install python-docx")

    # 确定输出路径
    if output_path is None:
        pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]
        output_dir = os.path.join(project_root, 'output', 'word')
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, f"{pdf_name}.docx")

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)

    # 创建临时目录
    temp_dir = tempfile.mkdtemp(prefix='pdf_to_word_')

    try:
        # Step 1: PDF 转图片（使用 JPEG 压缩）
        print("🖼️  正在将 PDF 转换为图片...")
        image_files = convert_pdf_to_images(pdf_path, temp_dir, dpi=dpi, quality=quality, use_jpeg=True)
        print(f"   ✅ 生成 {len(image_files)} 张图片")

        # Step 2: 创建 Word 文档
        print("📝 正在生成 Word 文档...")
        create_word_with_images(image_files, output_path)
        print(f"   ✅ Word 文档已保存")

        # 输出结果
        print()
        print("=" * 50)
        print("✅ 转换完成！")
        print("=" * 50)
        print(f"📄 Word 文件: {output_path}")
        print(f"📊 页数: {len(image_files)} 页")
        print(f"🎨 图片 DPI: {dpi}, JPEG 质量: {quality}")
        print()

        return {
            'success': True,
            'word_path': output_path,
            'page_count': len(image_files)
        }

    finally:
        # 清理临时文件
        if not keep_temp and os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
            print(f"🗑️  已清理临时文件")


# ============================================================================
# 命令行入口
# ============================================================================

def main():
    """命令行入口"""
    parser = argparse.ArgumentParser(description='PDF 转 Word 工具 - 将 PDF 每页转为图片插入 Word')
    parser.add_argument('pdf_path', help='PDF 文件路径')
    parser.add_argument('--output', '-o', help='输出 Word 文件路径（可选）')
    parser.add_argument('--dpi', '-d', type=int, default=100,
                        help='图片 DPI（默认 100，数值越大图片越清晰但文件越大）')
    parser.add_argument('--quality', '-q', type=int, default=70,
                        help='JPEG 质量 1-100（默认 70）')
    parser.add_argument('--keep-temp', action='store_true',
                        help='保留临时图片文件（用于调试）')

    args = parser.parse_args()

    try:
        result = pdf_to_word_with_images(
            pdf_path=args.pdf_path,
            output_path=args.output,
            dpi=args.dpi,
            quality=args.quality,
            keep_temp=args.keep_temp
        )
        return 0 if result['success'] else 1

    except Exception as e:
        print(f"❌ 转换失败: {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == '__main__':
    sys.exit(main())
